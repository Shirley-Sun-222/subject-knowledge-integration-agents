from __future__ import annotations

import io
import re
import shutil
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from typing import Callable

from ..config import settings
from ..utils.text import normalize_space, split_chapters, strip_repeated_headers


class ParseError(RuntimeError):
    pass


ParseProgress = Callable[[str, int, int], None]
TOP_LEVEL_CHAPTER_RE = re.compile(r"^(第\s*[一二三四五六七八九十百千万0-9IVXivx]+\s*章\b.*|Chapter\s+\d+\b.*)$", re.IGNORECASE)


def parse_textbook(path: Path, filename: str, progress: ParseProgress | None = None) -> dict:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        parsed = _parse_pdf(path, progress=progress)
    elif suffix in {".md", ".markdown", ".txt"}:
        if progress is not None:
            progress("reading_textbook", 1, 1)
        parsed = _parse_plain_text(path)
    elif suffix == ".docx":
        if progress is not None:
            progress("reading_textbook", 0, 1)
        parsed = _parse_docx(path)
        if progress is not None:
            progress("reading_textbook", 1, 1)
    else:
        raise ParseError(f"Unsupported file format: {suffix}")

    title = Path(filename).stem
    chapters = parsed.get("chapters") or split_chapters(parsed["text"], title=title, total_pages=parsed["total_pages"])
    return {
        "filename": filename,
        "title": title,
        "format": suffix.replace(".", ""),
        "total_pages": parsed["total_pages"],
        "total_chars": sum(chapter["char_count"] for chapter in chapters),
        "chapters": chapters,
    }


def _parse_plain_text(path: Path) -> dict:
    raw = path.read_text(encoding="utf-8", errors="ignore")
    return {"text": normalize_space(raw), "total_pages": 1}


def _parse_docx(path: Path) -> dict:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - optional dependency branch
        raise ParseError("python-docx is required to parse DOCX files") from exc

    document = Document(str(path))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
    return {"text": normalize_space(text), "total_pages": max(1, len(document.paragraphs) // 8)}


def _parse_pdf(path: Path, progress: ParseProgress | None = None) -> dict:
    try:
        import fitz
    except Exception as exc:  # pragma: no cover - optional dependency branch
        raise ParseError("PyMuPDF is required to parse PDF files") from exc

    document = fitz.open(path)
    total_pages = max(1, len(document))
    toc_spans = _extract_top_level_toc_spans(document, total_pages)
    mode = _detect_pdf_mode(document, progress)
    pages = _extract_page_texts(path, total_pages, progress)
    ocr_targets = _ocr_target_indices(mode, pages)
    if ocr_targets:
        ocr_pages = _extract_ocr_texts(path, ocr_targets, progress)
        for page_index, ocr_text in ocr_pages.items():
            if len(normalize_space(ocr_text)) > len(normalize_space(pages[page_index])):
                pages[page_index] = ocr_text
    text = normalize_space("\n\n".join(pages))
    if len(text) < 20:
        hint = "PDF has no extractable text"
        if settings.ocr_enabled:
            hint += "; OCR did not produce usable text. Check tesseract language data and OCR_MAX_PAGES."
        else:
            hint += "; enable OCR_ENABLED=1 for scanned PDFs."
        raise ParseError(hint)
    chapters = _chapters_from_toc_pages(toc_spans, pages, total_pages) if toc_spans else None
    return {"text": text, "total_pages": total_pages, "chapters": chapters}


def _detect_pdf_mode(document, progress: ParseProgress | None = None) -> str:
    sample_total = min(max(1, len(document)), min(settings.ocr_max_pages, 4))
    sparse_pages = 0
    for sample_index in range(sample_total):
        page_text = normalize_space(document[sample_index].get_text("text"))
        if len(page_text) < 20:
            sparse_pages += 1
        if progress is not None:
            progress("detecting_pdf_mode", sample_index + 1, sample_total)
    if sparse_pages == 0:
        return "digital"
    if sparse_pages == sample_total:
        return "scanned"
    return "mixed"


def _should_ocr_page(mode: str, page_text: str, page_index: int) -> bool:
    if not settings.ocr_enabled:
        return False
    normalized = normalize_space(page_text)
    if mode == "digital":
        return len(normalized) == 0
    return len(normalized) < 20


def _extract_page_texts(path: Path, total_pages: int, progress: ParseProgress | None = None) -> list[str]:
    indices = list(range(total_pages))
    workers = _resolve_pdf_workers(settings.pdf_text_extract_workers, total_pages)
    if workers == 1:
        pages = []
        for current, index in enumerate(indices, start=1):
            pages.append(_extract_single_page_text(path, index))
            if progress is not None:
                progress("reading_pdf_pages", current, total_pages)
        return pages

    page_map: dict[int, str] = {}
    completed = 0
    with ThreadPoolExecutor(max_workers=workers, thread_name_prefix="pdf-text") as executor:
        futures = {
            executor.submit(_extract_text_batch, path, batch): batch
            for batch in _page_batches(indices, workers)
        }
        for future in as_completed(futures):
            result = future.result()
            page_map.update(result)
            completed += len(result)
            if progress is not None:
                progress("reading_pdf_pages", completed, total_pages)
    return [page_map[index] for index in indices]


def _extract_ocr_texts(path: Path, indices: list[int], progress: ParseProgress | None = None) -> dict[int, str]:
    workers = _resolve_pdf_workers(settings.pdf_ocr_workers, len(indices))
    if workers == 1:
        pages = {}
        for current, index in enumerate(indices, start=1):
            pages[index] = _ocr_single_page(path, index)
            if progress is not None:
                progress("ocr_pdf_pages", current, len(indices))
        return pages

    pages: dict[int, str] = {}
    completed = 0
    with ThreadPoolExecutor(max_workers=workers, thread_name_prefix="pdf-ocr") as executor:
        futures = {
            executor.submit(_ocr_batch, path, batch): batch
            for batch in _page_batches(indices, workers)
        }
        for future in as_completed(futures):
            result = future.result()
            pages.update(result)
            completed += len(result)
            if progress is not None:
                progress("ocr_pdf_pages", completed, len(indices))
    return pages


def _extract_text_batch(path: Path, indices: list[int]) -> dict[int, str]:
    import fitz

    document = fitz.open(path)
    try:
        return {index: _extract_page_text_from_document(document, index) for index in indices}
    finally:
        document.close()


def _ocr_batch(path: Path, indices: list[int]) -> dict[int, str]:
    import fitz

    document = fitz.open(path)
    try:
        return {index: _ocr_pdf_page(document.load_page(index)) for index in indices}
    finally:
        document.close()


def _extract_single_page_text(path: Path, index: int) -> str:
    import fitz

    document = fitz.open(path)
    try:
        return _extract_page_text_from_document(document, index)
    finally:
        document.close()


def _ocr_single_page(path: Path, index: int) -> str:
    import fitz

    document = fitz.open(path)
    try:
        return _ocr_pdf_page(document.load_page(index))
    finally:
        document.close()


def _extract_page_text_from_document(document, index: int) -> str:
    lines = document.load_page(index).get_text("text").splitlines()
    return "\n".join(strip_repeated_headers(lines))


def _page_batches(indices: list[int], workers: int) -> list[list[int]]:
    if not indices:
        return []
    batch_size = max(1, len(indices) // workers)
    return [indices[offset : offset + batch_size] for offset in range(0, len(indices), batch_size)]


def _resolve_pdf_workers(configured: int, page_count: int) -> int:
    if page_count <= 1:
        return 1
    return max(1, min(configured, page_count))


def _ocr_target_indices(mode: str, pages: list[str]) -> list[int]:
    if not settings.ocr_enabled:
        return []
    if mode == "digital":
        return [index for index, text in enumerate(pages) if len(normalize_space(text)) == 0]
    if mode == "scanned":
        return list(range(len(pages)))
    return [index for index, text in enumerate(pages) if len(normalize_space(text)) < 20]


def _extract_top_level_toc_spans(document, total_pages: int) -> list[dict[str, int | str]]:
    toc = document.get_toc(simple=True)
    if not toc:
        return []

    matching = []
    for level, title, page in toc:
        normalized_title = _normalize_toc_title(title)
        if page <= 0 or page > total_pages:
            continue
        if TOP_LEVEL_CHAPTER_RE.match(normalized_title):
            matching.append((level, normalized_title, page))
    if not matching:
        return []

    levels = sorted({item[0] for item in matching})
    selected_level = next((level for level in levels if sum(1 for item in matching if item[0] == level) >= 3), levels[0])
    page_to_title: dict[int, str] = {}
    for level, title, page in matching:
        if level != selected_level:
            continue
        previous = page_to_title.get(page)
        if previous is None or len(title) > len(previous):
            page_to_title[page] = title

    ordered = sorted(page_to_title.items())
    spans: list[dict[str, int | str]] = []
    for index, (page_start, title) in enumerate(ordered):
        next_page = ordered[index + 1][0] if index + 1 < len(ordered) else total_pages + 1
        spans.append(
            {
                "title": title,
                "page_start": page_start,
                "page_end": max(page_start, next_page - 1),
            }
        )
    return spans


def _normalize_toc_title(title: str) -> str:
    normalized = normalize_space(title.replace("\u3000", " ").replace("\r", " ").replace("\n", " "))
    normalized = re.sub(r"\.{2,}", " ", normalized)
    normalized = re.sub(r"\s{2,}", " ", normalized)
    return normalized.strip()


def _chapters_from_toc_pages(toc_spans: list[dict[str, int | str]], pages: list[str], total_pages: int) -> list[dict[str, int | str]] | None:
    chapters = []
    for position, span in enumerate(toc_spans, start=1):
        page_start = int(span["page_start"])
        page_end = min(int(span["page_end"]), total_pages)
        chapter_text = normalize_space("\n\n".join(pages[page_start - 1 : page_end]))
        if len(chapter_text) < 20:
            continue
        chapters.append(
            {
                "title": str(span["title"]),
                "page_start": page_start,
                "page_end": page_end,
                "content": chapter_text,
                "char_count": len(chapter_text),
                "position": position,
            }
        )
    return chapters or None


def _ocr_pdf_page(page) -> str:
    try:
        import pytesseract
        from PIL import Image
        import fitz
    except Exception:
        return ""
    _configure_tesseract_cmd(pytesseract)
    try:
        # 2x render improves OCR accuracy without making sample uploads impractically slow.
        pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
        image = Image.open(io.BytesIO(pixmap.tobytes("png")))
    except Exception:
        return ""
    lang = _available_ocr_lang(settings.ocr_lang)
    try:
        return pytesseract.image_to_string(image, lang=lang).strip()
    except Exception:
        if lang != "eng":
            try:
                return pytesseract.image_to_string(image, lang="eng").strip()
            except Exception:
                return ""
        return ""


def _available_ocr_lang(requested: str) -> str:
    try:
        import pytesseract

        _configure_tesseract_cmd(pytesseract)
        available = set(pytesseract.get_languages(config=""))
    except Exception:
        return "eng"
    parts = [part for part in requested.split("+") if part]
    selected = [part for part in parts if part in available]
    if selected:
        return "+".join(selected)
    return "eng" if "eng" in available else (sorted(available)[0] if available else "eng")


def _configure_tesseract_cmd(pytesseract_module) -> None:
    command = shutil.which("tesseract")
    if command is None:
        for candidate in ["/opt/homebrew/bin/tesseract", "/usr/local/bin/tesseract", "/usr/bin/tesseract"]:
            if Path(candidate).exists():
                command = candidate
                break
    if command:
        pytesseract_module.pytesseract.tesseract_cmd = command
